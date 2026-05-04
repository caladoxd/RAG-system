import asyncio
import io
import logging
import os
import re
import zipfile
from pathlib import Path
from typing import Any, Final, Literal

import httpx
import tiktoken

from .embedding_service import embed_texts

logger = logging.getLogger(__name__)

TOKEN_LIMIT: Final[int] = 2000
OVERLAP_RATIO: Final[float] = 0.2
SEMANTIC_DISTANCE_PERCENTILE: Final[int] = 90
GENERATION_MODEL: Final[str] = os.getenv("GENERATION_MODEL", "qwen/qwen3-4b")
GENERATION_TIMEOUT_S: Final[float] = 120.0
DEFAULT_MAX_OUTPUT_TOKENS: Final[int] = int(os.getenv("GENERATION_MAX_TOKENS", "600"))
DEFAULT_TEMPERATURE: Final[float] = float(os.getenv("GENERATION_TEMPERATURE", "0.2"))
MAX_CONTEXT_CHARS_PER_CHUNK: Final[int] = 2000

_encoding: tiktoken.Encoding | None = None


def _get_encoding() -> tiktoken.Encoding:
    global _encoding
    if _encoding is None:
        try:
            _encoding = tiktoken.encoding_for_model("gpt-4o")
        except KeyError:
            _encoding = tiktoken.get_encoding("cl100k_base")
    return _encoding


def _count_tokens(text: str) -> int:
    return len(_get_encoding().encode(text))


def _cosine_similarity(a: list[float], b: list[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b, strict=True))
    na = sum(x * x for x in a) ** 0.5
    nb = sum(x * x for x in b) ** 0.5
    if na == 0 or nb == 0:
        return 0.0
    return dot / (na * nb)


def _percentile_threshold(values: list[float], p: int) -> float:
    if not values:
        return float("inf")
    s = sorted(values)
    k = max(0, min(len(s) - 1, int(round((p / 100.0) * (len(s) - 1)))))
    return s[k]


def _semantic_chunks_from_embeddings(sentences: list[str], embeddings: list[list[float]]) -> list[str]:
    n = len(sentences)
    if n == 0:
        return []
    if n == 1:
        return [sentences[0]]

    distances: list[float] = []
    for i in range(n - 1):
        sim = _cosine_similarity(embeddings[i], embeddings[i + 1])
        distances.append(1.0 - sim)

    threshold = _percentile_threshold(distances, SEMANTIC_DISTANCE_PERCENTILE)
    breaks: list[int] = [0]
    for i, d in enumerate(distances):
        if d >= threshold:
            breaks.append(i + 1)
    breaks.append(n)

    chunks: list[str] = []
    for a, b in zip(breaks, breaks[1:], strict=False):
        piece = " ".join(sentences[a:b]).strip()
        if piece:
            chunks.append(piece)
    return chunks if chunks else [" ".join(sentences)]


def _paragraph_semantic_fallback(document: str) -> list[str]:
    parts = [p.strip() for p in re.split(r"\n\n+", document.strip()) if p.strip()]
    return parts if parts else ([document.strip()] if document.strip() else [])


def _token_overlap_suffix(text: str, overlap_tokens: int) -> str:
    enc = _get_encoding()
    ids = enc.encode(text)
    if len(ids) <= overlap_tokens:
        return text
    return enc.decode(ids[-overlap_tokens:])


def _token_windows(text: str, max_tokens: int, overlap_tokens: int) -> list[str]:
    enc = _get_encoding()
    ids = enc.encode(text)
    if len(ids) <= max_tokens:
        return [text]
    chunks: list[str] = []
    start = 0
    while start < len(ids):
        end = min(start + max_tokens, len(ids))
        chunks.append(enc.decode(ids[start:end]))
        if end >= len(ids):
            break
        start = max(0, end - overlap_tokens)
    return chunks


def _paragraph_chunk(text: str, max_tokens: int, overlap_ratio: float) -> list[str]:
    enc = _get_encoding()
    overlap_tokens = max(1, int(max_tokens * overlap_ratio))
    paragraphs = [p.strip() for p in re.split(r"\n\n+", text.strip()) if p.strip()]
    if not paragraphs:
        return [text] if text.strip() else []

    chunks: list[str] = []
    buf: list[str] = []
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        ptoks = len(enc.encode(p))
        if ptoks > max_tokens:
            if buf:
                chunks.append("\n\n".join(buf))
                buf = []
            chunks.extend(_token_windows(p, max_tokens, overlap_tokens))
            i += 1
            continue

        candidate = "\n\n".join(buf + [p]) if buf else p
        if len(enc.encode(candidate)) <= max_tokens:
            buf.append(p)
            i += 1
            continue

        if buf:
            emitted = "\n\n".join(buf)
            chunks.append(emitted)
            suffix = _token_overlap_suffix(emitted, overlap_tokens)
            combined = f"{suffix}\n\n{p}" if suffix else p
            if len(enc.encode(combined)) <= max_tokens:
                buf = [suffix, p] if suffix else [p]
            else:
                chunks.extend(_token_windows(combined, max_tokens, overlap_tokens))
                buf = []
            i += 1
        else:
            buf.append(p)
            i += 1

    if buf:
        chunks.append("\n\n".join(buf))
    return chunks


def _is_docx_bytes(data: bytes) -> bool:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            return "word/document.xml" in zf.namelist()
    except zipfile.BadZipFile:
        return False


def _content_type_base(content_type: str | None) -> str:
    if not content_type:
        return ""
    return content_type.split(";")[0].strip().lower()


_TEXT_TYPES = frozenset(
    {
        "text/plain",
        "text/markdown",
        "text/csv",
        "text/xml",
        "application/json",
        "application/xml",
    }
)


def _looks_like_utf8_plaintext(data: bytes, max_scan: int = 262_144) -> bool:
    """True when bytes look like UTF-8 text (no filename / magic bytes)."""
    if not data:
        return False
    sample = data[: min(len(data), max_scan)]
    if b"\x00" in sample:
        return False
    try:
        s = sample.decode("utf-8")
    except UnicodeDecodeError:
        return False
    if not s:
        return True
    good = sum(1 for c in s if c.isprintable() or c in "\n\r\t\f\v")
    return good / len(s) >= 0.80


def _sniff_format(
    data: bytes, filename: str, content_type: str | None = None
) -> Literal["pdf", "docx", "txt"]:
    lower = filename.lower()
    if lower.endswith(".txt"):
        return "txt"
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".docx"):
        return "docx"
    ct = _content_type_base(content_type)
    if ct in _TEXT_TYPES or ct.startswith("text/"):
        return "txt"
    if len(data) >= 4 and data[:4] == b"%PDF":
        return "pdf"
    if _is_docx_bytes(data):
        return "docx"
    if _looks_like_utf8_plaintext(data):
        return "txt"
    raise ValueError(
        "Unsupported document type; expected PDF, DOCX, or UTF-8 text. "
        "Set filename query param or X-Filename (e.g. notes.txt), send Content-Type: text/plain, "
        "or use raw UTF-8 bytes without NUL characters."
    )


def _extract_pdf_sync(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        t = (page.extract_text() or "").strip()
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def _extract_docx_sync(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _extract_txt_sync(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def _build_rag_context(chunks: list[dict[str, Any]]) -> str:
    parts: list[str] = []
    for i, c in enumerate(chunks, start=1):
        doc_id = c.get("doc_id") or ""
        section_id = c.get("section_id") or ""
        chunk_index = c.get("chunk_index")
        text = str(c.get("text") or "").strip()
        if not text:
            continue
        if len(text) > MAX_CONTEXT_CHARS_PER_CHUNK:
            text = f"{text[:MAX_CONTEXT_CHARS_PER_CHUNK]}..."
        parts.append(
            f"[{i}] doc_id={doc_id} section_id={section_id} chunk_index={chunk_index}\n{text}"
        )
    return "\n\n".join(parts)    


async def extract_text(
    document: bytes | str,
    *,
    filename: str | None = None,
    content_type: str | None = None,
) -> str:
    """Return plain text from a PDF, DOCX, or TXT. Pass file bytes (and optional filename) or a path string."""
    if isinstance(document, str):
        path = Path(document)
        if not path.is_file():
            raise FileNotFoundError(str(path))
        data = path.read_bytes()
        name = filename or path.name
        ct = content_type
    else:
        data = document
        name = filename or ""
        ct = content_type
    if not data:
        return ""
    kind = _sniff_format(data, name, ct)
    if kind == "txt":
        return _extract_txt_sync(data)
    if kind == "pdf":
        return await asyncio.to_thread(_extract_pdf_sync, data)
    return await asyncio.to_thread(_extract_docx_sync, data)


async def chunk_document(document: str) -> list[str]:
    doc = document.strip()
    if not doc:
        return []

    sentences = [p.strip() for p in re.split(r"(?<=[.!?])\s+", doc) if p.strip()]
    if not sentences:
        return [doc]

    semantic_chunks: list[str]
    embeddings = await embed_texts(sentences)
    if embeddings is not None and len(embeddings) == len(sentences):
        semantic_chunks = _semantic_chunks_from_embeddings(sentences, embeddings)
    else:
        semantic_chunks = _paragraph_semantic_fallback(doc)

    if not semantic_chunks:
        semantic_chunks = [doc]

    final: list[str] = []
    for piece in semantic_chunks:
        if _count_tokens(piece) <= TOKEN_LIMIT:
            final.append(piece)
        else:
            final.extend(_paragraph_chunk(piece, TOKEN_LIMIT, OVERLAP_RATIO))
    return final


async def generate_answer(
    query: str,
    chunks: list[dict[str, Any]],
    *,
    temperature: float = DEFAULT_TEMPERATURE,
    max_tokens: int = DEFAULT_MAX_OUTPUT_TOKENS,
) -> str:
    q = query.strip()
    if not q:
        return ""
    context = _build_rag_context(chunks)
    if not context:
        return "I could not find relevant context to answer your question."

    base = os.getenv("OPENAI_BASE_URL", "http://localhost:1234/v1").rstrip("/")
    key = os.getenv("OPENAI_API_KEY", "lm-studio")
    url = f"{base}/chat/completions"

    system_prompt = (
        "You are a RAG assistant. Answer only using the provided context snippets. "
        "If context is insufficient, say that clearly. Keep the answer concise and factual."
    )
    user_prompt = (
        f"Question:\n{q}\n\n"
        f"Context snippets:\n{context}\n\n"
        "Provide the answer and include snippet citations like [1], [2] when used."
    )

    payload = {
        "model": GENERATION_MODEL,
        "temperature": temperature,
        "max_tokens": max_tokens,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    }
    timeout = httpx.Timeout(GENERATION_TIMEOUT_S, connect=10.0)
    try:
        async with httpx.AsyncClient(timeout=timeout) as client:
            response = await client.post(
                url,
                headers={"Authorization": f"Bearer {key}"},
                json=payload,
            )
            response.raise_for_status()
            data = response.json()
    except Exception as e:
        logger.warning("Generation request failed: %s", e)
        raise RuntimeError("Could not generate answer from retrieved context.") from e

    choices = data.get("choices") or []
    if not choices:
        raise RuntimeError("Generation returned no choices.")
    message = choices[0].get("message") or {}
    answer = (message.get("content") or "").strip()
    if not answer:
        raise RuntimeError("Generation returned an empty answer.")
    return answer

